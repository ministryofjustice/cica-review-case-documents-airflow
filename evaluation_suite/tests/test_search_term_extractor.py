"""Unit tests for search_term_extractor.py.

All tests build minimal in-memory Word documents with python-docx to avoid
any dependency on the real ``search_terms_source.docx`` file.
"""

import csv
import io
from pathlib import Path

import docx

from evaluation_suite.search_evaluation.multi_case.search_term_extractor import (
    BENCHMARK_TERMS,
    _split_cell_to_terms,
    extract_and_write_search_terms,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _add_case_row(table, case_number: int, question: str, col2: str, col3: str) -> None:
    """Append a 4-column row for *case_number* to *table*."""
    row = table.add_row()
    row.cells[0].text = f"Case {case_number}"
    row.cells[1].text = question
    row.cells[2].text = col2
    row.cells[3].text = col3


def _build_simple_doc(rows: list[tuple[int, str, str, str]]) -> docx.Document:
    """Build a Document with a single table populated from *rows*."""
    doc = docx.Document()
    if not rows:
        return doc
    table = doc.add_table(rows=0, cols=4)
    for case_number, question, col2, col3 in rows:
        _add_case_row(table, case_number, question, col2, col3)
    return doc


def _write_and_read(doc: docx.Document, tmp_path: Path) -> tuple[Path, Path]:
    """Save *doc* to *tmp_path* and return ``(docx_path, output_dir)``."""
    docx_path = tmp_path / "source.docx"
    buf = io.BytesIO()
    doc.save(buf)
    docx_path.write_bytes(buf.getvalue())
    output_dir = tmp_path / "cases"
    return docx_path, output_dir


def _read_csv_terms(csv_path: Path) -> list[str]:
    with open(csv_path, newline="", encoding="utf-8") as f:
        return [row["search_term"] for row in csv.DictReader(f)]


# ---------------------------------------------------------------------------
# _split_cell_to_terms
# ---------------------------------------------------------------------------


def test_split_cell_to_terms_splits_on_commas():
    assert _split_cell_to_terms("knee injury, back pain, wrist") == ["knee injury", "back pain", "wrist"]


def test_split_cell_to_terms_splits_on_newlines():
    result = _split_cell_to_terms("head trauma\nbrain injury")
    assert result == ["head trauma", "brain injury"]


def test_split_cell_to_terms_splits_on_full_stop():
    result = _split_cell_to_terms("skull fracture. Ongoing treatments")
    assert result == ["skull fracture", "Ongoing treatments"]


def test_split_cell_to_terms_splits_on_forward_slash():
    result = _split_cell_to_terms("Sexual assault/ abuse")
    assert "Sexual assault" in result
    assert "abuse" in result


def test_split_cell_to_terms_splits_on_trailing_hyphen_space():
    """'back injury- possible exacerbation' → two terms."""
    result = _split_cell_to_terms("back injury- possible exacerbation")
    assert "back injury" in result
    assert "possible exacerbation" in result


def test_split_cell_to_terms_splits_on_space_hyphen():
    """'care or supervision -answered yes' → two terms."""
    result = _split_cell_to_terms("care or supervision -answered yes")
    assert "care or supervision" in result


def test_split_cell_to_terms_preserves_compound_hyphen_words():
    """'pre-existing' must not be split — hyphen has no adjacent whitespace."""
    result = _split_cell_to_terms("pre-existing back injury")
    assert result == ["pre-existing back injury"]


def test_split_cell_to_terms_strips_ie_prefix():
    result = _split_cell_to_terms("ie mental health")
    assert result == ["mental health"]


def test_split_cell_to_terms_strips_eg_prefix():
    result = _split_cell_to_terms("eg acupuncture")
    assert result == ["acupuncture"]


def test_split_cell_to_terms_strips_leading_dashes():
    result = _split_cell_to_terms("- knee injury\n- back pain")
    assert result == ["knee injury", "back pain"]


def test_split_cell_to_terms_filters_short_strings():
    result = _split_cell_to_terms("of, a, knee injury")
    assert result == ["knee injury"]


def test_split_cell_to_terms_filters_long_terms():
    """Terms with more than 6 words are sentence fragments and must be dropped."""
    long_term = "pre existing therapy for mental health and medical records"
    result = _split_cell_to_terms(long_term)
    assert result == []


def test_split_cell_to_terms_keeps_six_word_terms():
    """Six words is the maximum allowed."""
    result = _split_cell_to_terms("dependency on medication for thirty years")
    assert result == ["dependency on medication for thirty years"]


def test_split_cell_to_terms_normalises_internal_whitespace():
    result = _split_cell_to_terms("knee   injury")
    assert result == ["knee injury"]


# ---------------------------------------------------------------------------
# extract_and_write_search_terms — basic extraction
# ---------------------------------------------------------------------------


def test_extracts_terms_from_tags_row(tmp_path):
    doc = _build_simple_doc(
        [
            (1, "What tags/categories apply?", "accident, fall", "injury, trauma"),
        ]
    )
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1])

    terms = _read_csv_terms(output_dir / "26-700001" / "search_terms.csv")
    assert "accident" in terms
    assert "fall" in terms
    assert "injury" in terms
    assert "trauma" in terms


def test_extracts_terms_from_contextual_row(tmp_path):
    doc = _build_simple_doc(
        [
            (2, "Contextual information", "hospital records", "medical notes"),
        ]
    )
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[2])

    terms = _read_csv_terms(output_dir / "26-700002" / "search_terms.csv")
    assert "hospital records" in terms
    assert "medical notes" in terms


def test_ignores_non_matching_question_rows(tmp_path):
    doc = _build_simple_doc(
        [
            (1, "What tags/categories apply?", "accident", "fall"),
            (1, "Elements to look for", "should_be_ignored", "also_ignored"),
        ]
    )
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1])

    terms = _read_csv_terms(output_dir / "26-700001" / "search_terms.csv")
    assert "should_be_ignored" not in terms
    assert "also_ignored" not in terms


def test_benchmark_terms_prepended_to_every_case(tmp_path):
    doc = _build_simple_doc(
        [
            (1, "Tags/categories", "knee injury", ""),
            (2, "Tags/categories", "brain trauma", ""),
        ]
    )
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1, 2])

    for case_ref in ("26-700001", "26-700002"):
        terms = _read_csv_terms(output_dir / case_ref / "search_terms.csv")
        for bench in BENCHMARK_TERMS:
            assert bench in terms, f"Benchmark term '{bench}' missing from {case_ref}"
        # Benchmark terms come first
        assert terms[: len(BENCHMARK_TERMS)] == BENCHMARK_TERMS


def test_case_zero_excluded_when_not_requested(tmp_path):
    """Case 0 is a training case and must not produce a CSV when excluded."""
    doc = _build_simple_doc(
        [
            (0, "Tags/categories", "training term", ""),
            (1, "Tags/categories", "real term", ""),
        ]
    )
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1])

    assert not (output_dir / "26-700000").exists()
    assert (output_dir / "26-700001" / "search_terms.csv").exists()


def test_output_directory_created_if_missing(tmp_path):
    doc = _build_simple_doc([(3, "Tags/categories", "sample term", "")])
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    output_dir_nested = output_dir / "nested" / "deep"
    extract_and_write_search_terms(docx_path, output_dir_nested, case_numbers=[3])

    assert (output_dir_nested / "26-700003" / "search_terms.csv").exists()


def test_combined_table_multiple_cases(tmp_path):
    """Table 0 in the real document mixes Cases 0-3 in a single table."""
    rows = [
        (1, "Tags/categories", "case1_term", ""),
        (2, "Tags/categories", "case2_term", ""),
        (3, "Tags/categories", "case3_term", ""),
    ]
    doc = _build_simple_doc(rows)
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1, 2, 3])

    assert "case1_term" in _read_csv_terms(output_dir / "26-700001" / "search_terms.csv")
    assert "case2_term" in _read_csv_terms(output_dir / "26-700002" / "search_terms.csv")
    assert "case3_term" in _read_csv_terms(output_dir / "26-700003" / "search_terms.csv")
    # Terms must not bleed across cases
    assert "case1_term" not in _read_csv_terms(output_dir / "26-700002" / "search_terms.csv")


def test_case_label_with_embedded_newlines(tmp_path):
    r"""Handles Col 0 text like 'Case 1\n\n\nCase Attribute' (real document format)."""
    doc = docx.Document()
    table = doc.add_table(rows=0, cols=4)
    row = table.add_row()
    row.cells[0].text = "Case 1\n\n\nCase Attribute"
    row.cells[1].text = "Tags/categories"
    row.cells[2].text = "embedded label term"
    row.cells[3].text = ""
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1])

    terms = _read_csv_terms(output_dir / "26-700001" / "search_terms.csv")
    assert "embedded label term" in terms


def test_case_specific_terms_deduplicated(tmp_path):
    """Duplicate terms extracted from the document are written only once."""
    doc = _build_simple_doc(
        [
            (1, "Tags/categories", "knee injury, knee injury", "Knee Injury"),
        ]
    )
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1])

    terms = _read_csv_terms(output_dir / "26-700001" / "search_terms.csv")
    assert terms.count("knee injury") == 1


def test_returns_term_counts_per_case(tmp_path):
    doc = _build_simple_doc(
        [
            (1, "Tags/categories", "term_a, term_b", "term_c"),
        ]
    )
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    result = extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1])

    expected_count = len(BENCHMARK_TERMS) + 3  # 3 specific terms
    assert result == {"26-700001": expected_count}


def test_total_terms_capped_at_max(tmp_path):
    """When extracted terms exceed _MAX_TERMS_PER_CASE the CSV is truncated."""
    from evaluation_suite.search_evaluation.multi_case.search_term_extractor import _MAX_TERMS_PER_CASE

    # Build enough specific terms to exceed the cap
    excess = _MAX_TERMS_PER_CASE - len(BENCHMARK_TERMS) + 10
    many_terms = ", ".join(f"term number {i:02d}" for i in range(excess))
    doc = _build_simple_doc([(1, "Tags/categories", many_terms, "")])
    docx_path, output_dir = _write_and_read(doc, tmp_path)

    result = extract_and_write_search_terms(docx_path, output_dir, case_numbers=[1])

    assert result["26-700001"] == _MAX_TERMS_PER_CASE
    terms = _read_csv_terms(output_dir / "26-700001" / "search_terms.csv")
    assert len(terms) == _MAX_TERMS_PER_CASE
    # Benchmark terms are always preserved
    assert terms[: len(BENCHMARK_TERMS)] == BENCHMARK_TERMS
